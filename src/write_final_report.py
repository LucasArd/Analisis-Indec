from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(".")
REPORT_PATH = ROOT / "informe.docx"
TABLES = ROOT / "outputs" / "tables"
FIGURES = ROOT / "outputs" / "figures"

SPANISH_REPLACEMENTS = {
    "Evolucion": "Evolución",
    "evolucion": "evolución",
    "Introduccion": "Introducción",
    "introduccion": "introducción",
    "metodologia": "metodología",
    "Metodologia": "Metodología",
    "analisis": "análisis",
    "Analisis": "Análisis",
    "comparacion": "comparación",
    "Comparacion": "Comparación",
    "seleccion": "selección",
    "Seleccion": "Selección",
    "dinamica": "dinámica",
    "Dinamica": "Dinámica",
    "energetica": "energética",
    "Energetica": "Energética",
    "informacion": "información",
    "Informacion": "Información",
    "inflacion": "inflación",
    "Inflacion": "Inflación",
    "poblacion": "población",
    "Poblacion": "Población",
    "ocupacion": "ocupación",
    "Ocupacion": "Ocupación",
    "desocupacion": "desocupación",
    "Desocupacion": "Desocupación",
    "educacion": "educación",
    "Educacion": "Educación",
    "regresion": "regresión",
    "Regresion": "Regresión",
    "imputacion": "imputación",
    "Imputacion": "Imputación",
    "estimacion": "estimación",
    "Estimacion": "Estimación",
    "dispersion": "dispersión",
    "Dispersion": "Dispersión",
    "distribucion": "distribución",
    "Distribucion": "Distribución",
    "posicion": "posición",
    "Posicion": "Posición",
    "medicion": "medición",
    "Medicion": "Medición",
    "codigo": "código",
    "Codigo": "Código",
    "codigos": "códigos",
    "Codigos": "Códigos",
    "periodo": "período",
    "Periodo": "Período",
    "trimestre": "trimestre",
    "graficos": "gráficos",
    "Graficos": "Gráficos",
    "grafico": "gráfico",
    "Grafico": "Gráfico",
    "estadistica": "estadística",
    "Estadistica": "Estadística",
    "estadisticas": "estadísticas",
    "Estadisticas": "Estadísticas",
    "economicamente": "económicamente",
    "Economicamente": "Económicamente",
    "ademas": "además",
    "Ademas": "Además",
    "tambien": "también",
    "Tambien": "También",
    "terminos": "términos",
    "Terminos": "Términos",
    "mas ": "más ",
    "Mas ": "Más ",
    "maximo": "máximo",
    "Maximo": "Máximo",
    "minimo": "mínimo",
    "Minimo": "Mínimo",
    "critico": "crítico",
    "Critico": "Crítico",
    "tecnica": "técnica",
    "Tecnica": "Técnica",
    "tecnico": "técnico",
    "Tecnico": "Técnico",
    "patagonica": "patagónica",
    "Patagonica": "Patagónica",
    "turismo": "turismo",
    "estacionalidad": "estacionalidad",
    "monetario": "monetario",
    "automaticamente": "automáticamente",
    "Automaticamente": "Automáticamente",
    "intercuartilico": "intercuartílico",
    "Intercuartilico": "Intercuartílico",
    "asimetrica": "asimétrica",
    "Asimetrica": "Asimétrica",
    "asimetria": "asimetría",
    "Asimetria": "Asimetría",
    "heterogeneas": "heterogéneas",
    "Heterogeneas": "Heterogéneas",
    "puntuales": "puntuales",
    "disponibles": "disponibles",
    "lineal": "lineal",
    "logaritmo": "logaritmo",
    "variacion": "variación",
    "Variacion": "Variación",
    "explicativa": "explicativa",
    "sistematicas": "sistemáticas",
    "Sistematicas": "Sistemáticas",
    "categorias": "categorías",
    "Categorias": "Categorías",
    "anio": "año",
    "Anio": "Año",
    "Neuquen": "Neuquén",
    "sociodemograficas": "sociodemográficas",
    "Sociodemograficas": "Sociodemográficas",
    "caida": "caída",
    "Caida": "Caída",
    "esta disponible": "está disponible",
    "esta atravesado": "está atravesado",
    "urbaños": "urbanos",
    "engañosa": "engañosa",
    "enganosas": "engañosas",
    "proximos": "próximos",
    "Proximos": "Próximos",
    "Sintesis": "Síntesis",
    "sintesis": "síntesis",
    "evaluacion": "evaluación",
    "Evaluacion": "Evaluación",
    "multiple": "múltiple",
    "Multiple": "Múltiple",
    "unicamente": "únicamente",
    "Unicamente": "Únicamente",
    "relacion": "relación",
    "Relacion": "Relación",
    "especificacion": "especificación",
    "Especificacion": "Especificación",
    "regularizacion": "regularización",
    "Regularizacion": "Regularización",
    "entreno": "entrenó",
    "Entreno": "Entrenó",
    "evaluo": "evaluó",
    "Evaluo": "Evaluó",
    "particion": "partición",
    "Particion": "Partición",
    "desvio": "desvío",
    "Desvio": "Desvío",
    "practicamente": "prácticamente",
    "Practicamente": "Prácticamente",
    "geografica": "geográfica",
    "Geografica": "Geográfica",
    "historico": "histórico",
    "Historico": "Histórico",
    "unica": "única",
    "Unica": "Única",
    "ocupacion": "ocupación",
    "Descripcion": "Descripción",
    "descripcion": "descripción",
    "Calculo": "Cálculo",
    "calculo": "cálculo",
    "univariada": "univariada",
    "Atipicos": "Atípicos",
    "atipicos": "atípicos",
    "Energeticas": "Energéticas",
    "energeticas": "energéticas",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fmt_pct(value) -> str:
    return f"{float(value):.1f}%"


def fmt_money(value) -> str:
    return f"${float(value):,.0f}".replace(",", ".")


def normalize_spanish(text: str) -> str:
    for source, target in SPANISH_REPLACEMENTS.items():
        text = text.replace(source, target)
    return text


def normalize_doc_text(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath(".//w:drawing"):
                continue
            run.text = normalize_spanish(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xpath(".//w:drawing"):
                            continue
                        run.text = normalize_spanish(run.text)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(normalize_spanish(text))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(10.5)


def add_table(doc: Document, dataframe: pd.DataFrame, title: str) -> None:
    doc.add_paragraph(normalize_spanish(title), style="Caption")
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0].cells
    for i, col in enumerate(dataframe.columns):
        set_cell_text(header[i], normalize_spanish(str(col)), bold=True)
        set_cell_shading(header[i], "D9EAF7")

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], normalize_spanish(str(value)))


def add_figure(doc: Document, filename: str, caption: str, width: float = 5.8) -> None:
    path = FIGURES / filename
    doc.add_picture(str(path), width=Inches(width))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(normalize_spanish(caption), style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report() -> None:
    indicadores = pd.read_csv(TABLES / "indicadores_laborales_trimestrales.csv")
    ingresos = pd.read_csv(TABLES / "ingresos_reales_trimestrales.csv")
    resumen = pd.read_csv(TABLES / "resumen_inicial_final.csv")
    univariado = pd.read_csv(TABLES / "resumen_univariado.csv")
    lineales = pd.read_csv(TABLES / "modelos_lineales_metricas.csv")
    lineal_simple = pd.read_csv(TABLES / "regresion_lineal_simple_coeficientes.csv")
    lineal_multiple = pd.read_csv(TABLES / "regresion_lineal_multiple_coeficientes_top.csv")
    modelo = pd.read_csv(TABLES / "modelo_imputacion_metricas.csv")
    coef = pd.read_csv(TABLES / "modelo_imputacion_coeficientes_top.csv")
    subgrupos = pd.read_csv(TABLES / "indicadores_por_subgrupos.csv")
    imputados = pd.read_csv(TABLES / "ingresos_reales_imputados_trimestrales.csv")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)

    title = doc.add_heading(
        "Evolución del mercado laboral e ingresos reales", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Análisis comparativo de Neuquén-Plottier y Mar del Plata, EPH 2016T2-2025T4"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Fuente: elaboración propia con microdatos EPH-INDEC e IPC nacional.")

    doc.add_heading("1. Introducción", level=1)
    add_paragraph(
        doc,
        "El objetivo del trabajo es analizar la evolucion de la tasa de actividad, "
        "la tasa de empleo, la tasa de desocupacion y los ingresos de la poblacion "
        "en dos aglomerados urbanos relevados por la Encuesta Permanente de Hogares "
        "(EPH). La comparacion se concentra en Neuquen-Plottier, identificado con el "
        "código de aglomerado 17, y Mar del Plata, identificado con el código 34. "
        "La seleccion permite contrastar dos mercados laborales con perfiles distintos: "
        "uno asociado a la dinámica patagónica y energética, y otro con fuerte presencia "
        "de servicios, turismo y estacionalidad.",
    )
    add_paragraph(
        doc,
        "El período analizado cubre desde el segundo trimestre de 2016 hasta el cuarto "
        "trimestre de 2025, que es el tramo disponible en las bases usuarias descargadas "
        "desde INDEC. Para que los ingresos sean comparables en el tiempo, se los ajusta "
        "por inflacion mediante el IPC Nivel General Nacional, base diciembre de 2016. "
        "Los valores monetarios se expresan a precios del cuarto trimestre de 2025.",
    )

    doc.add_heading("2. Datos y metodología", level=1)
    add_paragraph(
        doc,
        "Se utilizan los microdatos de personas de la EPH. Las tasas laborales se "
        "calcularon con el ponderador PONDERA. La poblacion economicamente activa "
        "(PEA) se define como la suma ponderada de ocupados y desocupados. La tasa "
        "de actividad es PEA sobre poblacion total ponderada; la tasa de empleo es "
        "ocupados sobre poblacion total ponderada; y la tasa de desocupacion es "
        "desocupados sobre PEA. Para ingresos laborales se usa P21, ingreso de la "
        "ocupacion principal, y se pondera con PONDIIO cuando esta disponible.",
    )
    add_paragraph(
        doc,
        "Además del análisis agregado por trimestre y aglomerado, se incorporan variables "
        "sociodemograficas y laborales: sexo, edad, nivel educativo, rama de actividad "
        "(PP04B_COD) y ocupacion (PP04D_COD). Estas variables permiten pasar de una "
        "descripción puramente univariada a una lectura multivariada de la evolución "
        "de los indicadores.",
    )
    add_paragraph(
        doc,
        "El procesamiento fue realizado en Python mediante pandas, numpy, matplotlib, "
        "seaborn y scikit-learn. La eleccion del lenguaje no modifica la metodologia "
        "estadistica aplicada: las operaciones realizadas equivalen a procedimientos "
        "habituales de R, como agrupar y resumir datos, unir tablas, construir graficos "
        "y estimar modelos de regresion. Por ejemplo, los agrupamientos de pandas cumplen "
        "el mismo rol que group_by/summarise en dplyr, los cruces de tablas equivalen a "
        "joins y los graficos generados con seaborn/matplotlib cumplen una funcion similar "
        "a los realizados con ggplot2.",
    )

    summary_table = resumen.copy()
    summary_table = summary_table[
        [
            "aglomerado",
            "actividad_inicial",
            "actividad_final",
            "empleo_inicial",
            "empleo_final",
            "desocupacion_inicial",
            "desocupacion_final",
            "ingreso_mediano_real_inicial",
            "ingreso_mediano_real_final",
        ]
    ]
    summary_table.columns = [
        "Aglomerado",
        "Actividad inicial",
        "Actividad final",
        "Empleo inicial",
        "Empleo final",
        "Desoc. inicial",
        "Desoc. final",
        "Ingreso mediano inicial",
        "Ingreso mediano final",
    ]
    for col in summary_table.columns[1:7]:
        summary_table[col] = summary_table[col].map(fmt_pct)
    for col in ["Ingreso mediano inicial", "Ingreso mediano final"]:
        summary_table[col] = summary_table[col].map(fmt_money)
    add_table(doc, summary_table, "Tabla 1. Resumen inicial-final de indicadores seleccionados")

    doc.add_heading("3. Evolución de actividad, empleo y desocupación", level=1)
    add_paragraph(
        doc,
        "La evolucion de las tasas laborales muestra diferencias relevantes entre los "
        "aglomerados. Mar del Plata parte en 2016T2 con una tasa de actividad de 45,8% "
        "y llega a 52,2% en 2025T4. Su tasa de empleo tambien aumenta, de 40,5% a 47,2%. "
        "Sin embargo, la desocupacion continua siendo relativamente elevada al final del "
        "periodo: 9,5% en 2025T4. Esto sugiere una recuperacion de la participacion y "
        "del empleo, aunque con persistencia de presion sobre el mercado de trabajo.",
    )
    add_paragraph(
        doc,
        "Neuquen-Plottier muestra una trayectoria distinta. La tasa de actividad se "
        "mantiene alrededor de niveles similares entre 2016T2 y 2025T4, pasando de "
        "46,2% a 45,7%. La tasa de empleo, en cambio, sube de 42,2% a 44,7%, mientras "
        "que la desocupacion cae de 8,6% a 2,3%. En terminos comparativos, el cierre "
        "del periodo muestra un mercado laboral mas ajustado en Neuquen-Plottier que "
        "en Mar del Plata.",
    )
    add_figure(doc, "tasa_actividad.png", "Figura 1. Tasa de actividad por trimestre")
    add_figure(doc, "tasa_empleo.png", "Figura 2. Tasa de empleo por trimestre")
    add_figure(doc, "tasa_desocupacion.png", "Figura 3. Tasa de desocupacion por trimestre")
    add_paragraph(
        doc,
        "El shock de 2020 aparece con claridad en ambas series. En 2020T2, Mar del Plata "
        "alcanza una desocupacion de 26,0%, mientras que Neuquen-Plottier llega a 13,5%. "
        "La diferencia de magnitudes refuerza la importancia de comparar aglomerados: "
        "los efectos de una misma coyuntura macroeconomica no se traducen de forma "
        "identica en todos los mercados laborales urbanos.",
    )

    doc.add_heading("4. Ingresos reales y medidas de posición", level=1)
    add_paragraph(
        doc,
        "Para analizar ingresos se trabaja con valores reales, deflactados por IPC. "
        "Dado que los ingresos suelen tener distribuciones asimetricas, el analisis "
        "prioriza la mediana y los percentiles 25, 75 y 90, ademas de la media. Esto "
        "permite interpretar no solo el nivel promedio, sino tambien la posicion de "
        "distintos tramos de la distribucion.",
    )

    income_sel = ingresos[ingresos["PERIODO"].isin(["2016T4", "2020T2", "2025T4"])][
        [
            "PERIODO",
            "AGLOMERADO_NOMBRE",
            "ingreso_medio_real",
            "ingreso_mediano_real",
            "p25_real",
            "p75_real",
            "p90_real",
        ]
    ].copy()
    income_sel.columns = ["Periodo", "Aglomerado", "Media", "Mediana", "P25", "P75", "P90"]
    for col in ["Media", "Mediana", "P25", "P75", "P90"]:
        income_sel[col] = income_sel[col].map(fmt_money)
    add_table(doc, income_sel, "Tabla 2. Ingresos reales de ocupados, medidas seleccionadas")
    add_figure(doc, "ingreso_mediano_real.png", "Figura 4. Ingreso mediano real de ocupados")
    add_figure(doc, "boxplot_ingresos_reales_2025.png", "Figura 5. Distribucion de ingresos reales en 2025")
    add_paragraph(
        doc,
        "En 2025T4, la mediana del ingreso real de la ocupacion principal es de "
        "$1.000.000 en Mar del Plata y de $1.300.000 en Neuquen-Plottier. La diferencia "
        "tambien se observa en los tramos superiores: el percentil 90 alcanza $2.200.000 "
        "en Mar del Plata y $3.000.000 en Neuquen-Plottier. Estos resultados son "
        "compatibles con una estructura de ingresos mas alta y dispersa en Neuquen-Plottier.",
    )

    doc.add_heading("5. Exploración univariada, no respuesta y valores atípicos", level=1)
    add_paragraph(
        doc,
        "La exploracion univariada permite evaluar la calidad inicial de las variables. "
        "Para cada variable relevante se calcularon casos, faltantes, media, mediana, "
        "percentiles extremos y valores atipicos mediante el criterio del rango "
        "intercuartilico. En ingresos, la presencia de ceros, valores negativos y "
        "valores extremos obliga a tomar decisiones explicitas antes de estimar medidas "
        "de tendencia central o ajustar modelos.",
    )
    uni = univariado.copy()
    if "variable_nombre" not in uni.columns:
        uni["variable_nombre"] = uni["variable"]
    uni = uni[
        [
            "variable",
            "variable_nombre",
            "faltantes_pct",
            "mediana",
            "p01",
            "p99",
            "atipicos_iqr",
        ]
    ]
    uni.columns = [
        "Variable",
        "Nombre",
        "Faltantes %",
        "Mediana",
        "P01",
        "P99",
        "Atípicos IQR",
    ]
    uni["Faltantes %"] = uni["Faltantes %"].map(lambda x: f"{x:.1f}%")
    for col in ["Mediana", "P01", "P99"]:
        uni[col] = uni[col].map(lambda x: f"{x:.1f}")
    add_table(doc, uni, "Tabla 3. Exploracion univariada de variables seleccionadas")
    add_paragraph(
        doc,
        "En P21 se observa una mediana igual a cero cuando se considera toda la poblacion, "
        "lo cual es esperable porque incluye personas no ocupadas o sin ingreso laboral. "
        "Por eso el analisis de ingresos se restringe a ocupados y a ingresos positivos "
        "para el calculo de medidas reales. Para la imputacion se considera no respuesta "
        "operativa cuando P21 es faltante o negativo; los ceros no se imputan "
        "automaticamente porque pueden representar ausencia efectiva de ingreso laboral.",
    )

    doc.add_heading("6. Análisis multivariado por sexo, edad y educación", level=1)
    add_paragraph(
        doc,
        "El analisis por subgrupos muestra que la evolucion agregada no afecta por igual "
        "a toda la poblacion. La tasa de desocupacion por sexo permite observar brechas "
        "de insercion laboral, mientras que la tasa de empleo por nivel educativo muestra "
        "la relacion positiva entre credenciales educativas y participacion efectiva en "
        "el empleo. Estas aperturas complementan la comparacion entre aglomerados y "
        "evitan interpretar los promedios como situaciones homogeneas.",
    )
    add_figure(doc, "desocupacion_por_sexo.png", "Figura 6. Tasa de desocupacion por sexo")
    add_figure(doc, "empleo_por_nivel_educativo_2025.png", "Figura 7. Tasa de empleo por nivel educativo")

    latest_edu = subgrupos[
        (subgrupos["variable"] == "NIVEL_ED_DESC")
        & (subgrupos["anio"] == 2025)
        & (~subgrupos["categoria"].isin(["Ns/Nr", "Sin dato"]))
    ][["aglomerado", "categoria", "tasa_empleo", "tasa_desocupacion"]].copy()
    latest_edu.columns = ["Aglomerado", "Nivel educativo", "Tasa empleo", "Tasa desocupacion"]
    latest_edu["Tasa empleo"] = latest_edu["Tasa empleo"].map(fmt_pct)
    latest_edu["Tasa desocupacion"] = latest_edu["Tasa desocupacion"].map(fmt_pct)
    add_table(doc, latest_edu, "Tabla 4. Indicadores por nivel educativo, 2025")
    add_paragraph(
        doc,
        "Ademas de sexo y educacion, el modelo de ingresos incorpora edad, rama de "
        "actividad y ocupacion. Las variables PP04B_COD y PP04D_COD permiten capturar "
        "diferencias vinculadas al tipo de insercion laboral. Aunque en el informe se "
        "presentan de forma resumida, quedan disponibles en las tablas de coeficientes "
        "para revisar actividades u ocupaciones puntuales.",
    )

    doc.add_heading("7. Modelo de imputación de no respuesta de ingresos", level=1)
    add_paragraph(
        doc,
        "Como paso previo al modelo de imputacion se estimaron dos modelos de regresion "
        "lineal sobre el logaritmo del ingreso real de la ocupacion principal. El primero "
        "es una regresion lineal simple, que explica el ingreso unicamente a partir de la "
        "edad. El segundo es una regresion lineal multiple, que incorpora edad, sexo, "
        "nivel educativo, aglomerado, anio, trimestre, rama de actividad y ocupacion. "
        "Esta comparacion permite mostrar el aporte del analisis multivariado: una sola "
        "variable explica muy poco, mientras que el conjunto de caracteristicas personales "
        "y laborales captura mejor las diferencias de ingresos.",
    )
    linear_table = lineales.copy()
    linear_table["modelo"] = linear_table["modelo"].replace(
        {
            "Regresion lineal simple: log(P21 real) ~ edad": "Lineal simple: edad",
            "Regresion lineal multiple: log(P21 real) ~ edad + sexo + educacion + aglomerado + periodo + rama + ocupacion": "Lineal multiple",
        }
    )
    linear_table = linear_table[
        ["modelo", "n_train", "n_test", "mae_pesos_2025t4", "rmse_pesos_2025t4", "r2_log"]
    ]
    linear_table.columns = ["Modelo", "Train", "Test", "MAE", "RMSE", "R2 log"]
    linear_table["MAE"] = linear_table["MAE"].map(fmt_money)
    linear_table["RMSE"] = linear_table["RMSE"].map(fmt_money)
    linear_table["R2 log"] = linear_table["R2 log"].map(lambda x: f"{x:.3f}")
    add_table(doc, linear_table, "Tabla 5. Comparación de regresión lineal simple y múltiple")
    simple_effect = lineal_simple.loc[0, "impacto_pct"]
    add_paragraph(
        doc,
        f"En la regresion lineal simple, la edad estandarizada tiene un coeficiente "
        f"positivo: un aumento de un desvio estandar en la edad se asocia con un cambio "
        f"aproximado de {simple_effect:.1f}% en el ingreso real esperado. Sin embargo, "
        f"el R2 es practicamente nulo, por lo que la edad por si sola no alcanza para "
        f"explicar la heterogeneidad de ingresos.",
    )
    add_figure(
        doc,
        "regresion_simple_edad_log_ingreso.png",
        "Figura 8. Regresión lineal simple entre edad y logaritmo del ingreso real",
    )
    if "variable_nombre" not in lineal_multiple.columns:
        lineal_multiple["variable_nombre"] = lineal_multiple["variable"]
    multiple_small = lineal_multiple.head(6)[
        ["variable", "variable_nombre", "impacto_pct"]
    ].copy()
    multiple_small.columns = [
        "Variable",
        "Nombre",
        "Impacto porcentual aproximado",
    ]
    multiple_small["Impacto porcentual aproximado"] = multiple_small[
        "Impacto porcentual aproximado"
    ].map(lambda x: f"{x:.1f}%")
    add_table(doc, multiple_small, "Tabla 6. Efectos principales en la regresión lineal múltiple")
    add_paragraph(
        doc,
        "La regresion lineal multiple mejora claramente el ajuste. Esto indica que los "
        "ingresos no dependen de una unica dimension individual, sino de una combinacion "
        "de edad, sexo, credenciales educativas, ubicacion geografica, periodo historico "
        "y caracteristicas del empleo. En particular, las categorias de ocupacion y rama "
        "aparecen entre los efectos mas grandes, lo que confirma que el tipo de puesto "
        "es una variable central para explicar diferencias salariales.",
    )
    add_figure(
        doc,
        "regresion_multiple_residuos.png",
        "Figura 9. Diagnóstico de regresión múltiple: residuos frente a valores predichos",
    )
    add_paragraph(
        doc,
        "El gráfico de residuos permite evaluar visualmente si el modelo deja patrones "
        "sistemáticos sin explicar. Si los residuos se distribuyen de forma aproximadamente "
        "aleatoria alrededor de cero, el ajuste lineal resulta más razonable. En cambio, "
        "patrones marcados o cambios fuertes en la dispersión indicarían problemas de "
        "especificación u homocedasticidad.",
    )
    add_paragraph(
        doc,
        "Para cumplir el objetivo de aprobacion directa se desarrollo, ademas, un modelo "
        "Ridge para imputar la no respuesta de ingresos. Ridge mantiene una estructura "
        "lineal similar a la regresion multiple, pero agrega regularizacion para reducir "
        "la inestabilidad cuando hay muchas categorias de rama y ocupacion. Se utilizaron "
        "como variables independientes la edad, sexo, nivel educativo, aglomerado, anio, "
        "trimestre, rama de actividad y ocupacion. El modelo se entreno con ocupados de "
        "14 a 80 años con ingreso positivo y se evaluo mediante una particion train/test.",
    )
    m = modelo.iloc[0]
    model_table = pd.DataFrame(
        [
            {
                "Modelo": m["modelo"],
                "Train": int(m["n_train"]),
                "Test": int(m["n_test"]),
                "MAE": fmt_money(m["mae_pesos_2025t4"]),
                "RMSE": fmt_money(m["rmse_pesos_2025t4"]),
                "R2 log": f"{m['r2_log']:.3f}",
            }
        ]
    )
    add_table(doc, model_table, "Tabla 7. Evaluacion del modelo de imputacion")
    add_paragraph(
        doc,
        "El R2 sobre el logaritmo del ingreso es 0,460. Esto indica una capacidad "
        "explicativa moderada: el modelo captura diferencias sistematicas por edad, "
        "educacion, aglomerado, rama y ocupacion, pero deja una parte importante de la "
        "variacion individual sin explicar. El MAE y el RMSE se expresan en pesos reales "
        "a precios de 2025T4, por lo que deben interpretarse en el contexto de una "
        "distribucion de ingresos muy dispersa.",
    )

    if "variable_nombre" not in coef.columns:
        coef["variable_nombre"] = coef["variable"]
    coef_small = coef.head(8)[["variable", "variable_nombre", "impacto_pct"]].copy()
    coef_small.columns = ["Variable", "Nombre", "Impacto porcentual aproximado"]
    coef_small["Impacto porcentual aproximado"] = coef_small[
        "Impacto porcentual aproximado"
    ].map(lambda x: f"{x:.1f}%")
    add_table(doc, coef_small, "Tabla 8. Variables con mayor efecto estimado en el modelo Ridge")
    add_paragraph(
        doc,
        "La interpretacion de los coeficientes se realiza sobre el logaritmo del ingreso: "
        "un coeficiente positivo implica un ingreso esperado mayor, manteniendo constantes "
        "las demas variables; uno negativo implica menor ingreso esperado. Entre los efectos "
        "mas grandes aparecen categorias de ocupacion y rama, lo que confirma que las "
        "caracteristicas del puesto son centrales para explicar diferencias salariales. "
        "El uso de Ridge reduce la inestabilidad de los coeficientes cuando hay muchas "
        "categorias y algunas tienen pocos casos.",
    )

    imp_last = imputados[imputados["PERIODO"].eq("2025T4")][
        [
            "AGLOMERADO_NOMBRE",
            "no_respuesta_p21",
            "ocupados",
            "ingreso_mediano_real_original",
            "ingreso_mediano_real_imputado",
        ]
    ].copy()
    imp_last.columns = [
        "Aglomerado",
        "No respuesta P21",
        "Ocupados",
        "Mediana original",
        "Mediana imputada",
    ]
    for col in ["Mediana original", "Mediana imputada"]:
        imp_last[col] = imp_last[col].map(fmt_money)
    add_table(doc, imp_last, "Tabla 9. Ingresos originales e imputados, 2025T4")

    doc.add_heading("8. Conclusiones", level=1)
    add_paragraph(
        doc,
        "La comparacion muestra dos trayectorias laborales diferenciadas. Mar del Plata "
        "termina el periodo con mayor actividad y empleo que al inicio, pero tambien con "
        "una desocupacion final relativamente alta. Neuquen-Plottier, en cambio, exhibe "
        "una fuerte reduccion de la desocupacion y una mejora de la tasa de empleo, aun "
        "sin aumento de la tasa de actividad. En ingresos reales, Neuquen-Plottier presenta "
        "niveles medianos superiores a Mar del Plata al final del periodo.",
    )
    add_paragraph(
        doc,
        "El ajuste por inflacion resulta indispensable: sin deflactar, la evolucion nominal "
        "de los ingresos seria poco informativa en un periodo de alta inflacion. Las medidas "
        "de posicion muestran que las diferencias entre aglomerados no se reducen al promedio, "
        "sino que tambien aparecen en la mediana y en percentiles altos. Finalmente, el modelo "
        "de imputacion aporta una herramienta estadistica para tratar la no respuesta de "
        "ingresos, aunque sus resultados deben interpretarse como aproximaciones y no como "
        "sustitutos perfectos del dato observado.",
    )
    add_paragraph(
        doc,
        "Como limitaciones, el analisis utiliza IPC nacional, por lo que no capta diferencias "
        "regionales de precios entre Patagonia y region pampeana. Ademas, la imputacion depende "
        "de las variables observadas disponibles en la EPH y no captura factores no relevados "
        "que tambien inciden sobre los ingresos. Aun asi, el trabajo permite interpretar la "
        "evolucion historica de los principales indicadores laborales y de ingresos para los "
        "dos aglomerados seleccionados.",
    )

    normalize_doc_text(doc)
    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
