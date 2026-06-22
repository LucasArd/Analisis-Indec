from __future__ import annotations

from io import StringIO
from pathlib import Path
from urllib.request import Request, urlopen

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches
from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LinearRegression, Ridge
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler


ROOT = Path(".")
INDIVIDUAL_PATH = ROOT / "data" / "infoProcesada" / "eph" / "eph_individual_2016T2_2025T4.csv.gz"
INFLATION_DIR = ROOT / "data" / "infoInflacion"
OUTPUT_TABLES = ROOT / "outputs" / "tables"
OUTPUT_FIGURES = ROOT / "outputs" / "figures"
REPORT_DIR = ROOT / "informe"

AGGLOMERATES = {
    17: "Neuquen-Plottier",
    34: "Mar del Plata",
}

VARIABLE_DESCRIPTIONS = {
    "P21": "Ingreso de la ocupación principal",
    "ITF": "Ingreso total familiar",
    "IPCF": "Ingreso per cápita familiar",
    "CH06": "Edad",
    "CH04": "Sexo",
    "NIVEL_ED": "Nivel educativo",
    "NIVEL_ED_DESC": "Nivel educativo",
    "PP04B_COD": "Rama de actividad",
    "PP04D_COD": "Ocupación",
    "SEXO": "Sexo",
    "GRUPO_EDAD": "Grupo de edad",
    "AGLOMERADO": "Aglomerado",
    "ANO4": "Año",
    "TRIMESTRE": "Trimestre",
    "Edad estandarizada": "Edad",
}

USECOLS = [
    "ANO4",
    "TRIMESTRE",
    "AGLOMERADO",
    "PONDERA",
    "PONDIIO",
    "ESTADO",
    "CH04",
    "CH06",
    "NIVEL_ED",
    "P21",
    "ITF",
    "IPCF",
    "PP04B_COD",
    "PP04D_COD",
]

IPC_SERIES_ID = "148.3_INIVELNAL_DICI_M_26"
IPC_URL = (
    "https://apis.datos.gob.ar/series/api/series/"
    f"?ids={IPC_SERIES_ID}&format=csv&limit=5000"
)


def ensure_dirs() -> None:
    INFLATION_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_TABLES.mkdir(parents=True, exist_ok=True)
    OUTPUT_FIGURES.mkdir(parents=True, exist_ok=True)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)


def describe_variable(variable: str) -> str:
    variable = str(variable)
    if variable in VARIABLE_DESCRIPTIONS:
        return VARIABLE_DESCRIPTIONS[variable]
    if "PP04D_COD" in variable:
        return "Ocupación"
    if "PP04B_COD" in variable:
        return "Rama de actividad"
    if "NIVEL_ED" in variable:
        return "Nivel educativo"
    if "CH04" in variable:
        return "Sexo"
    if "CH06" in variable:
        return "Edad"
    if "AGLOMERADO" in variable:
        return "Aglomerado"
    if "ANO4" in variable:
        return "Año"
    if "TRIMESTRE" in variable:
        return "Trimestre"
    return "Variable derivada del modelo"


def download_ipc() -> pd.DataFrame:
    request = Request(IPC_URL, headers={"User-Agent": "Mozilla/5.0"})
    with urlopen(request, timeout=60) as response:
        content = response.read().decode("utf-8")

    ipc = pd.read_csv(StringIO(content))
    ipc.columns = ["fecha", "ipc"]
    ipc["fecha"] = pd.to_datetime(ipc["fecha"])
    ipc["anio"] = ipc["fecha"].dt.year
    ipc["trimestre"] = ipc["fecha"].dt.quarter
    ipc.to_csv(INFLATION_DIR / "ipc_nacional_mensual.csv", index=False)

    ipc_q = (
        ipc.groupby(["anio", "trimestre"], as_index=False)
        .agg(ipc=("ipc", "mean"))
        .sort_values(["anio", "trimestre"])
    )
    ipc_q.to_csv(INFLATION_DIR / "ipc_nacional_trimestral.csv", index=False)
    return ipc_q


def load_filtered_data() -> pd.DataFrame:
    frames = []
    for chunk in pd.read_csv(
        INDIVIDUAL_PATH,
        usecols=USECOLS,
        chunksize=100_000,
        low_memory=False,
    ):
        chunk = chunk[chunk["AGLOMERADO"].isin(AGGLOMERATES)]
        frames.append(chunk)

    df = pd.concat(frames, ignore_index=True)
    df.columns = df.columns.str.upper()

    numeric_cols = [
        "ANO4",
        "TRIMESTRE",
        "AGLOMERADO",
        "PONDERA",
        "PONDIIO",
        "ESTADO",
        "CH04",
        "CH06",
        "NIVEL_ED",
        "P21",
        "ITF",
        "IPCF",
        "PP04B_COD",
        "PP04D_COD",
    ]
    for col in numeric_cols:
        df[col] = pd.to_numeric(df[col], errors="coerce")

    df["AGLOMERADO_NOMBRE"] = df["AGLOMERADO"].map(AGGLOMERATES)
    df["PERIODO"] = df["ANO4"].astype("Int64").astype(str) + "T" + df["TRIMESTRE"].astype("Int64").astype(str)
    df["FECHA"] = pd.PeriodIndex.from_fields(
        year=df["ANO4"].astype(int),
        quarter=df["TRIMESTRE"].astype(int),
        freq="Q",
    ).to_timestamp()
    df["SEXO"] = df["CH04"].map({1: "Varon", 2: "Mujer"}).fillna("Sin dato")
    df["GRUPO_EDAD"] = pd.cut(
        df["CH06"],
        bins=[-np.inf, 13, 24, 34, 49, 64, np.inf],
        labels=["0-13", "14-24", "25-34", "35-49", "50-64", "65+"],
    )
    df["NIVEL_ED_DESC"] = df["NIVEL_ED"].map(
        {
            1: "Primaria incompleta",
            2: "Primaria completa",
            3: "Secundaria incompleta",
            4: "Secundaria completa",
            5: "Superior/univ. incompleta",
            6: "Superior/univ. completa",
            7: "Sin instruccion",
            9: "Ns/Nr",
        }
    ).fillna("Sin dato")
    return df


def weighted_sum(df: pd.DataFrame, mask: pd.Series, weight: str) -> float:
    return df.loc[mask, weight].fillna(0).sum()


def weighted_quantile(values: pd.Series, weights: pd.Series, quantile: float) -> float:
    valid = values.notna() & weights.notna() & (weights > 0)
    values = values[valid].to_numpy()
    weights = weights[valid].to_numpy()
    if len(values) == 0:
        return np.nan
    sorter = np.argsort(values)
    values = values[sorter]
    weights = weights[sorter]
    cumulative = np.cumsum(weights)
    cutoff = quantile * weights.sum()
    return values[np.searchsorted(cumulative, cutoff, side="left")]


def calculate_labor_indicators(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    group_cols = ["ANO4", "TRIMESTRE", "FECHA", "PERIODO", "AGLOMERADO", "AGLOMERADO_NOMBRE"]
    for keys, group in df.groupby(group_cols, observed=True):
        total = weighted_sum(group, group["ESTADO"].isin([1, 2, 3, 4]), "PONDERA")
        pea = weighted_sum(group, group["ESTADO"].isin([1, 2]), "PONDERA")
        occupied = weighted_sum(group, group["ESTADO"].eq(1), "PONDERA")
        unemployed = weighted_sum(group, group["ESTADO"].eq(2), "PONDERA")
        rows.append(
            {
                **dict(zip(group_cols, keys)),
                "poblacion_ponderada": total,
                "pea_ponderada": pea,
                "ocupados_ponderados": occupied,
                "desocupados_ponderados": unemployed,
                "tasa_actividad": pea / total * 100 if total else np.nan,
                "tasa_empleo": occupied / total * 100 if total else np.nan,
                "tasa_desocupacion": unemployed / pea * 100 if pea else np.nan,
                "casos": len(group),
            }
        )
    out = pd.DataFrame(rows).sort_values(["AGLOMERADO", "ANO4", "TRIMESTRE"])
    out.to_csv(OUTPUT_TABLES / "indicadores_laborales_trimestrales.csv", index=False)
    return out


def calculate_subgroup_tables(df: pd.DataFrame) -> None:
    rows = []
    for subgroup in ["SEXO", "GRUPO_EDAD", "NIVEL_ED_DESC"]:
        group_cols = ["ANO4", "AGLOMERADO_NOMBRE", subgroup]
        for keys, group in df.groupby(group_cols, observed=True):
            total = weighted_sum(group, group["ESTADO"].isin([1, 2, 3, 4]), "PONDERA")
            pea = weighted_sum(group, group["ESTADO"].isin([1, 2]), "PONDERA")
            occupied = weighted_sum(group, group["ESTADO"].eq(1), "PONDERA")
            unemployed = weighted_sum(group, group["ESTADO"].eq(2), "PONDERA")
            rows.append(
                {
                    "variable": subgroup,
                    "variable_nombre": describe_variable(subgroup),
                    "categoria": keys[2],
                    "anio": keys[0],
                    "aglomerado": keys[1],
                    "tasa_actividad": pea / total * 100 if total else np.nan,
                    "tasa_empleo": occupied / total * 100 if total else np.nan,
                    "tasa_desocupacion": unemployed / pea * 100 if pea else np.nan,
                    "casos": len(group),
                }
            )
    pd.DataFrame(rows).to_csv(OUTPUT_TABLES / "indicadores_por_subgrupos.csv", index=False)


def add_real_income(df: pd.DataFrame, ipc_q: pd.DataFrame) -> tuple[pd.DataFrame, float]:
    ipc_q = ipc_q.rename(columns={"anio": "ANO4", "trimestre": "TRIMESTRE", "ipc": "IPC"})
    base_ipc = float(
        ipc_q.loc[(ipc_q["ANO4"] == 2025) & (ipc_q["TRIMESTRE"] == 4), "IPC"].iloc[0]
    )
    df = df.merge(ipc_q, on=["ANO4", "TRIMESTRE"], how="left")
    df["P21_REAL"] = np.where(df["P21"] > 0, df["P21"] * base_ipc / df["IPC"], np.nan)
    df["ITF_REAL"] = np.where(df["ITF"] > 0, df["ITF"] * base_ipc / df["IPC"], np.nan)
    df["IPCF_REAL"] = np.where(df["IPCF"] > 0, df["IPCF"] * base_ipc / df["IPC"], np.nan)
    return df, base_ipc


def calculate_income_tables(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    group_cols = ["ANO4", "TRIMESTRE", "FECHA", "PERIODO", "AGLOMERADO_NOMBRE"]
    workers = df[df["ESTADO"].eq(1)].copy()
    workers["PESO_INGRESO"] = workers["PONDIIO"].fillna(workers["PONDERA"])

    for keys, group in workers.groupby(group_cols, observed=True):
        valid_income = group["P21_REAL"].notna()
        rows.append(
            {
                **dict(zip(group_cols, keys)),
                "ingreso_medio_real": np.average(
                    group.loc[valid_income, "P21_REAL"],
                    weights=group.loc[valid_income, "PESO_INGRESO"],
                )
                if valid_income.any()
                else np.nan,
                "ingreso_mediano_real": weighted_quantile(
                    group["P21_REAL"], group["PESO_INGRESO"], 0.50
                ),
                "p25_real": weighted_quantile(group["P21_REAL"], group["PESO_INGRESO"], 0.25),
                "p75_real": weighted_quantile(group["P21_REAL"], group["PESO_INGRESO"], 0.75),
                "p90_real": weighted_quantile(group["P21_REAL"], group["PESO_INGRESO"], 0.90),
                "ocupados": len(group),
                "ocupados_con_ingreso_valido": int(valid_income.sum()),
                "no_respuesta_ingreso": int((group["P21"].isna() | (group["P21"] < 0)).sum()),
                "tasa_no_respuesta_ingreso": (group["P21"].isna() | (group["P21"] < 0)).mean() * 100,
            }
        )
    out = pd.DataFrame(rows).sort_values(["AGLOMERADO_NOMBRE", "ANO4", "TRIMESTRE"])
    out.to_csv(OUTPUT_TABLES / "ingresos_reales_trimestrales.csv", index=False)
    return out


def calculate_univariate_summary(df: pd.DataFrame) -> pd.DataFrame:
    vars_to_check = ["P21", "ITF", "IPCF", "CH06", "NIVEL_ED", "PP04B_COD", "PP04D_COD"]
    rows = []
    for var in vars_to_check:
        values = df[var]
        numeric = pd.to_numeric(values, errors="coerce")
        q1 = numeric.quantile(0.25)
        q3 = numeric.quantile(0.75)
        iqr = q3 - q1
        rows.append(
            {
                "variable": var,
                "variable_nombre": describe_variable(var),
                "casos": len(values),
                "faltantes": int(values.isna().sum()),
                "faltantes_pct": values.isna().mean() * 100,
                "media": numeric.mean(),
                "mediana": numeric.median(),
                "p01": numeric.quantile(0.01),
                "p99": numeric.quantile(0.99),
                "atipicos_iqr": int(((numeric < q1 - 1.5 * iqr) | (numeric > q3 + 1.5 * iqr)).sum())
                if pd.notna(iqr)
                else 0,
            }
        )
    out = pd.DataFrame(rows)
    out.to_csv(OUTPUT_TABLES / "resumen_univariado.csv", index=False)
    return out


def build_income_model_data(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.Series]:
    model_df = df[
        (df["ESTADO"].eq(1))
        & (df["P21_REAL"].notna())
        & (df["P21_REAL"] > 0)
        & (df["CH06"].between(14, 80))
    ].copy()
    y = np.log(model_df["P21_REAL"])
    return model_df, y


def evaluate_log_income_model(
    model_name: str,
    model: Pipeline,
    X_train: pd.DataFrame,
    X_test: pd.DataFrame,
    y_train: pd.Series,
    y_test: pd.Series,
) -> dict:
    model.fit(X_train, y_train)
    pred_log = model.predict(X_test)
    pred_income = np.exp(pred_log)
    true_income = np.exp(y_test)
    return {
        "modelo": model_name,
        "n_train": len(X_train),
        "n_test": len(X_test),
        "mae_pesos_2025t4": mean_absolute_error(true_income, pred_income),
        "rmse_pesos_2025t4": mean_squared_error(true_income, pred_income) ** 0.5,
        "r2_log": r2_score(y_test, pred_log),
    }


def fit_linear_income_models(df: pd.DataFrame) -> pd.DataFrame:
    model_df, y = build_income_model_data(df)

    simple_features = ["CH06"]
    multiple_features = [
        "CH06",
        "CH04",
        "NIVEL_ED",
        "AGLOMERADO",
        "ANO4",
        "TRIMESTRE",
        "PP04B_COD",
        "PP04D_COD",
    ]

    X_train_s, X_test_s, y_train_s, y_test_s = train_test_split(
        model_df[simple_features], y, test_size=0.25, random_state=42
    )
    simple_model = Pipeline(
        steps=[
            (
                "preprocessor",
                ColumnTransformer(
                    transformers=[
                        (
                            "num",
                            Pipeline(
                                [
                                    ("imputer", SimpleImputer(strategy="median")),
                                    ("scaler", StandardScaler()),
                                ]
                            ),
                            simple_features,
                        )
                    ]
                ),
            ),
            ("regressor", LinearRegression()),
        ]
    )

    X_train_m, X_test_m, y_train_m, y_test_m = train_test_split(
        model_df[multiple_features], y, test_size=0.25, random_state=42
    )
    multiple_model = Pipeline(
        steps=[
            (
                "preprocessor",
                ColumnTransformer(
                    transformers=[
                        (
                            "num",
                            Pipeline(
                                [
                                    ("imputer", SimpleImputer(strategy="median")),
                                    ("scaler", StandardScaler()),
                                ]
                            ),
                            ["CH06", "ANO4", "TRIMESTRE"],
                        ),
                        (
                            "cat",
                            Pipeline(
                                [
                                    ("imputer", SimpleImputer(strategy="most_frequent")),
                                    ("onehot", OneHotEncoder(handle_unknown="ignore")),
                                ]
                            ),
                            ["CH04", "NIVEL_ED", "AGLOMERADO", "PP04B_COD", "PP04D_COD"],
                        ),
                    ]
                ),
            ),
            ("regressor", LinearRegression()),
        ]
    )

    metrics = pd.DataFrame(
        [
            evaluate_log_income_model(
                "Regresion lineal simple: log(P21 real) ~ edad",
                simple_model,
                X_train_s,
                X_test_s,
                y_train_s,
                y_test_s,
            ),
            evaluate_log_income_model(
                "Regresion lineal multiple: log(P21 real) ~ edad + sexo + educacion + aglomerado + periodo + rama + ocupacion",
                multiple_model,
                X_train_m,
                X_test_m,
                y_train_m,
                y_test_m,
            ),
        ]
    )

    simple_pred_test = simple_model.predict(X_test_s)
    multiple_pred_test = multiple_model.predict(X_test_m)

    simple_plot = X_test_s.copy()
    simple_plot["log_ingreso_real"] = y_test_s
    if len(simple_plot) > 4000:
        simple_plot = simple_plot.sample(4000, random_state=42)
    age_grid = pd.DataFrame({"CH06": np.linspace(14, 80, 100)})
    age_pred = simple_model.predict(age_grid)

    sns.set_theme(style="whitegrid")
    plt.figure(figsize=(9, 5))
    sns.scatterplot(
        data=simple_plot,
        x="CH06",
        y="log_ingreso_real",
        s=14,
        alpha=0.25,
        edgecolor=None,
    )
    plt.plot(age_grid["CH06"], age_pred, color="#C44E52", linewidth=2.5)
    plt.title("Regresion lineal simple: edad y log ingreso real")
    plt.xlabel("Edad")
    plt.ylabel("log(P21 real)")
    plt.tight_layout()
    plt.savefig(OUTPUT_FIGURES / "regresion_simple_edad_log_ingreso.png", dpi=160)
    plt.close()

    residuals = y_test_m - multiple_pred_test
    residual_plot = pd.DataFrame(
        {
            "log_ingreso_predicho": multiple_pred_test,
            "residuo": residuals,
        }
    )
    if len(residual_plot) > 5000:
        residual_plot = residual_plot.sample(5000, random_state=42)
    plt.figure(figsize=(9, 5))
    sns.scatterplot(
        data=residual_plot,
        x="log_ingreso_predicho",
        y="residuo",
        s=14,
        alpha=0.25,
        edgecolor=None,
    )
    plt.axhline(0, color="#C44E52", linewidth=2)
    plt.title("Regresion lineal multiple: residuos vs valores predichos")
    plt.xlabel("log(P21 real) predicho")
    plt.ylabel("Residuo")
    plt.tight_layout()
    plt.savefig(OUTPUT_FIGURES / "regresion_multiple_residuos.png", dpi=160)
    plt.close()

    simple_coef = simple_model.named_steps["regressor"].coef_[0]
    simple_intercept = simple_model.named_steps["regressor"].intercept_
    simple_coef_table = pd.DataFrame(
        [
            {
                "modelo": "Regresion lineal simple",
                "variable": "Edad estandarizada",
                "variable_nombre": describe_variable("Edad estandarizada"),
                "coef_log": simple_coef,
                "impacto_pct": (np.exp(simple_coef) - 1) * 100,
                "intercepto_log": simple_intercept,
            }
        ]
    )

    feature_names = multiple_model.named_steps["preprocessor"].get_feature_names_out()
    multiple_coefs = multiple_model.named_steps["regressor"].coef_
    multiple_coef_table = (
        pd.DataFrame({"variable": feature_names, "coef_log": multiple_coefs})
        .assign(variable_nombre=lambda x: x["variable"].map(describe_variable))
        .assign(impacto_pct=lambda x: (np.exp(x["coef_log"]) - 1) * 100)
        .sort_values("coef_log", key=lambda s: s.abs(), ascending=False)
        .head(30)
    )

    metrics.to_csv(OUTPUT_TABLES / "modelos_lineales_metricas.csv", index=False)
    simple_coef_table.to_csv(
        OUTPUT_TABLES / "regresion_lineal_simple_coeficientes.csv", index=False
    )
    multiple_coef_table.to_csv(
        OUTPUT_TABLES / "regresion_lineal_multiple_coeficientes_top.csv", index=False
    )
    return metrics


def fit_income_model(df: pd.DataFrame) -> tuple[pd.DataFrame, Pipeline, list[str]]:
    model_df, y = build_income_model_data(df)

    features = [
        "CH06",
        "CH04",
        "NIVEL_ED",
        "AGLOMERADO",
        "ANO4",
        "TRIMESTRE",
        "PP04B_COD",
        "PP04D_COD",
    ]
    X = model_df[features]

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.25, random_state=42
    )

    numeric_features = ["CH06", "ANO4", "TRIMESTRE"]
    categorical_features = ["CH04", "NIVEL_ED", "AGLOMERADO", "PP04B_COD", "PP04D_COD"]

    preprocessor = ColumnTransformer(
        transformers=[
            (
                "num",
                Pipeline(
                    [
                        ("imputer", SimpleImputer(strategy="median")),
                        ("scaler", StandardScaler()),
                    ]
                ),
                numeric_features,
            ),
            (
                "cat",
                Pipeline(
                    [
                        ("imputer", SimpleImputer(strategy="most_frequent")),
                        ("onehot", OneHotEncoder(handle_unknown="ignore")),
                    ]
                ),
                categorical_features,
            ),
        ]
    )

    model = Pipeline(
        steps=[
            ("preprocessor", preprocessor),
            ("regressor", Ridge(alpha=1.0)),
        ]
    )
    model.fit(X_train, y_train)
    pred_log = model.predict(X_test)
    pred_income = np.exp(pred_log)
    true_income = np.exp(y_test)

    metrics = pd.DataFrame(
        [
            {
                "modelo": "Ridge sobre log(P21 real)",
                "n_train": len(X_train),
                "n_test": len(X_test),
                "mae_pesos_2025t4": mean_absolute_error(true_income, pred_income),
                "rmse_pesos_2025t4": mean_squared_error(true_income, pred_income) ** 0.5,
                "r2_log": r2_score(y_test, pred_log),
            }
        ]
    )

    # Coefficients are interpreted on log income. For readability, keep largest effects.
    feature_names = model.named_steps["preprocessor"].get_feature_names_out()
    coefs = model.named_steps["regressor"].coef_
    coef_table = (
        pd.DataFrame({"variable": feature_names, "coef_log": coefs})
        .assign(variable_nombre=lambda x: x["variable"].map(describe_variable))
        .assign(impacto_pct=lambda x: (np.exp(x["coef_log"]) - 1) * 100)
        .sort_values("coef_log", key=lambda s: s.abs(), ascending=False)
        .head(30)
    )

    metrics.to_csv(OUTPUT_TABLES / "modelo_imputacion_metricas.csv", index=False)
    coef_table.to_csv(OUTPUT_TABLES / "modelo_imputacion_coeficientes_top.csv", index=False)
    return metrics, model, features


def calculate_imputed_income_tables(
    df: pd.DataFrame, model: Pipeline, features: list[str]
) -> pd.DataFrame:
    workers = df[df["ESTADO"].eq(1)].copy()
    workers["PESO_INGRESO"] = workers["PONDIIO"].fillna(workers["PONDERA"])
    workers["NO_RESPUESTA_P21"] = workers["P21"].isna() | (workers["P21"] < 0)
    workers["P21_REAL_IMPUTADO"] = workers["P21_REAL"]

    to_impute = workers["NO_RESPUESTA_P21"] & workers["IPC"].notna() & workers["CH06"].between(14, 80)
    if to_impute.any():
        pred_log = model.predict(workers.loc[to_impute, features])
        workers.loc[to_impute, "P21_REAL_IMPUTADO"] = np.exp(pred_log)

    rows = []
    group_cols = ["ANO4", "TRIMESTRE", "FECHA", "PERIODO", "AGLOMERADO_NOMBRE"]
    for keys, group in workers.groupby(group_cols, observed=True):
        valid_original = group["P21_REAL"].notna()
        valid_imputed = group["P21_REAL_IMPUTADO"].notna()
        rows.append(
            {
                **dict(zip(group_cols, keys)),
                "no_respuesta_p21": int(group["NO_RESPUESTA_P21"].sum()),
                "ocupados": len(group),
                "ingreso_mediano_real_original": weighted_quantile(
                    group.loc[valid_original, "P21_REAL"],
                    group.loc[valid_original, "PESO_INGRESO"],
                    0.50,
                ),
                "ingreso_mediano_real_imputado": weighted_quantile(
                    group.loc[valid_imputed, "P21_REAL_IMPUTADO"],
                    group.loc[valid_imputed, "PESO_INGRESO"],
                    0.50,
                ),
            }
        )

    out = pd.DataFrame(rows)
    out.to_csv(OUTPUT_TABLES / "ingresos_reales_imputados_trimestrales.csv", index=False)

    imputed_cases = workers[to_impute].copy()
    cols = [
        "ANO4",
        "TRIMESTRE",
        "AGLOMERADO_NOMBRE",
        "CH04",
        "CH06",
        "NIVEL_ED",
        "PP04B_COD",
        "PP04D_COD",
        "P21",
        "P21_REAL_IMPUTADO",
    ]
    imputed_cases[cols].to_csv(OUTPUT_TABLES / "casos_ingreso_imputado.csv", index=False)
    return out


def make_plots(indicators: pd.DataFrame, incomes: pd.DataFrame, df: pd.DataFrame) -> None:
    sns.set_theme(style="whitegrid")

    plot_specs = [
        ("tasa_actividad", "Tasa de actividad (%)", "tasa_actividad.png"),
        ("tasa_empleo", "Tasa de empleo (%)", "tasa_empleo.png"),
        ("tasa_desocupacion", "Tasa de desocupacion (%)", "tasa_desocupacion.png"),
    ]
    for y, ylabel, filename in plot_specs:
        plt.figure(figsize=(10, 5))
        sns.lineplot(data=indicators, x="FECHA", y=y, hue="AGLOMERADO_NOMBRE", marker="o")
        plt.title(ylabel + " por trimestre")
        plt.xlabel("Periodo")
        plt.ylabel(ylabel)
        plt.tight_layout()
        plt.savefig(OUTPUT_FIGURES / filename, dpi=160)
        plt.close()

    plt.figure(figsize=(10, 5))
    sns.lineplot(
        data=incomes,
        x="FECHA",
        y="ingreso_mediano_real",
        hue="AGLOMERADO_NOMBRE",
        marker="o",
    )
    plt.title("Ingreso mediano real de ocupados")
    plt.xlabel("Periodo")
    plt.ylabel("Pesos a precios de 2025T4")
    plt.tight_layout()
    plt.savefig(OUTPUT_FIGURES / "ingreso_mediano_real.png", dpi=160)
    plt.close()

    latest_year = int(df["ANO4"].max())
    latest = df[(df["ANO4"] == latest_year) & (df["ESTADO"].eq(1)) & df["P21_REAL"].notna()]
    latest = latest[latest["P21_REAL"] <= latest["P21_REAL"].quantile(0.99)]
    plt.figure(figsize=(9, 5))
    sns.boxplot(data=latest, x="AGLOMERADO_NOMBRE", y="P21_REAL")
    plt.title(f"Distribucion de ingresos reales de ocupados, {latest_year}")
    plt.xlabel("")
    plt.ylabel("Pesos a precios de 2025T4")
    plt.tight_layout()
    plt.savefig(OUTPUT_FIGURES / "boxplot_ingresos_reales_2025.png", dpi=160)
    plt.close()

    subgroup = pd.read_csv(OUTPUT_TABLES / "indicadores_por_subgrupos.csv")
    sex = subgroup[(subgroup["variable"] == "SEXO") & (subgroup["anio"] >= 2017)]
    plt.figure(figsize=(10, 5))
    sns.lineplot(
        data=sex,
        x="anio",
        y="tasa_desocupacion",
        hue="aglomerado",
        style="categoria",
        marker="o",
    )
    plt.title("Tasa de desocupacion por sexo")
    plt.xlabel("Anio")
    plt.ylabel("%")
    plt.tight_layout()
    plt.savefig(OUTPUT_FIGURES / "desocupacion_por_sexo.png", dpi=160)
    plt.close()

    edu = subgroup[
        (subgroup["variable"] == "NIVEL_ED_DESC")
        & (subgroup["anio"] == latest_year)
        & (~subgroup["categoria"].isin(["Ns/Nr", "Sin dato"]))
    ]
    plt.figure(figsize=(11, 5))
    sns.barplot(data=edu, x="categoria", y="tasa_empleo", hue="aglomerado")
    plt.title(f"Tasa de empleo por nivel educativo, {latest_year}")
    plt.xlabel("")
    plt.ylabel("%")
    plt.xticks(rotation=30, ha="right")
    plt.tight_layout()
    plt.savefig(OUTPUT_FIGURES / "empleo_por_nivel_educativo_2025.png", dpi=160)
    plt.close()


def first_last_summary(indicators: pd.DataFrame, incomes: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for aglo, group in indicators.groupby("AGLOMERADO_NOMBRE"):
        first = group.sort_values("FECHA").iloc[0]
        last = group.sort_values("FECHA").iloc[-1]
        income_group = incomes[incomes["AGLOMERADO_NOMBRE"] == aglo].dropna(
            subset=["ingreso_mediano_real"]
        )
        income_first = income_group.sort_values("FECHA").iloc[0]
        income_last = income_group.sort_values("FECHA").iloc[-1]
        rows.append(
            {
                "aglomerado": aglo,
                "periodo_inicial_laboral": first["PERIODO"],
                "periodo_final_laboral": last["PERIODO"],
                "actividad_inicial": first["tasa_actividad"],
                "actividad_final": last["tasa_actividad"],
                "empleo_inicial": first["tasa_empleo"],
                "empleo_final": last["tasa_empleo"],
                "desocupacion_inicial": first["tasa_desocupacion"],
                "desocupacion_final": last["tasa_desocupacion"],
                "periodo_inicial_ingreso": income_first["PERIODO"],
                "periodo_final_ingreso": income_last["PERIODO"],
                "ingreso_mediano_real_inicial": income_first["ingreso_mediano_real"],
                "ingreso_mediano_real_final": income_last["ingreso_mediano_real"],
            }
        )
    out = pd.DataFrame(rows)
    out.to_csv(OUTPUT_TABLES / "resumen_inicial_final.csv", index=False)
    return out


def write_report(
    indicators: pd.DataFrame,
    incomes: pd.DataFrame,
    univariate: pd.DataFrame,
    linear_metrics: pd.DataFrame,
    model_metrics: pd.DataFrame,
    summary: pd.DataFrame,
    imputed_incomes: pd.DataFrame,
) -> None:
    latest_period = indicators.sort_values("FECHA")["PERIODO"].iloc[-1]
    first_period = indicators.sort_values("FECHA")["PERIODO"].iloc[0]
    income_first_period = incomes.dropna(subset=["ingreso_mediano_real"]).sort_values("FECHA")[
        "PERIODO"
    ].iloc[0]

    md = f"""# Evolucion del mercado laboral e ingresos: Neuquen-Plottier vs Mar del Plata

## 1. Introduccion

Este informe analiza la evolucion de la tasa de actividad, la tasa de empleo, la tasa de desocupacion y los ingresos de la poblacion en dos aglomerados urbanos relevados por la EPH: Neuquen-Plottier y Mar del Plata. El periodo cubierto por las bases descargadas es {first_period}-{latest_period}. La comparacion combina una lectura temporal de los indicadores laborales con un analisis de ingresos reales, variables sociodemograficas y un modelo inicial de imputacion de ingresos.

## 2. Datos y metodologia

La fuente principal son los microdatos de la Encuesta Permanente de Hogares (EPH) publicados por INDEC. Se utilizan las bases de personas para el periodo 2016T2-2025T4. Los codigos de aglomerado empleados son 17 para Neuquen-Plottier y 34 para Mar del Plata. Las tasas laborales se calculan con el ponderador `PONDERA`. Para los ingresos de ocupados se usa `P21`, ponderado con `PONDIIO` cuando esta disponible.

La tasa de actividad se calcula como PEA sobre poblacion total ponderada; la tasa de empleo como ocupados sobre poblacion total ponderada; y la tasa de desocupacion como desocupados sobre PEA. Para ingresos reales se usa el IPC Nivel General Nacional, base diciembre 2016, publicado por INDEC/datos.gob.ar. Como la serie nacional mensual comienza en diciembre de 2016, el analisis de ingresos reales comienza efectivamente en {income_first_period}. Los valores se expresan a precios de 2025T4.

## 3. Evolucion de los indicadores laborales

![Tasa de actividad](../outputs/figures/tasa_actividad.png)

![Tasa de empleo](../outputs/figures/tasa_empleo.png)

![Tasa de desocupacion](../outputs/figures/tasa_desocupacion.png)

La comparacion muestra trayectorias laborales con oscilaciones importantes entre 2016 y 2025. El periodo 2020 se destaca por el impacto de la pandemia, visible en la caida de observaciones y en movimientos fuertes de actividad y empleo. La lectura comparativa debe hacerse mirando tanto niveles como cambios: Mar del Plata suele mostrar una estructura laboral sensible a estacionalidad y servicios, mientras que Neuquen-Plottier esta atravesado por el peso de actividades energeticas y dinamicas regionales patagonicas.

## 4. Ingresos reales

![Ingreso mediano real](../outputs/figures/ingreso_mediano_real.png)

![Distribucion de ingresos](../outputs/figures/boxplot_ingresos_reales_2025.png)

El ingreso mediano real permite observar la capacidad de compra de los ocupados una vez descontada la inflacion. Se prioriza la mediana por sobre la media porque los ingresos presentan alta asimetria y valores extremos. El boxplot de 2025 muestra que la dispersion de ingresos es considerable en ambos aglomerados, por lo que la comparacion no debe limitarse al promedio.

## 5. Analisis por subgrupos

![Desocupacion por sexo](../outputs/figures/desocupacion_por_sexo.png)

![Empleo por nivel educativo](../outputs/figures/empleo_por_nivel_educativo_2025.png)

El analisis por sexo y nivel educativo permite ver heterogeneidades que quedan ocultas en las tasas agregadas. En general, la insercion laboral tiende a mejorar con el nivel educativo, mientras que la desocupacion por sexo puede mostrar brechas persistentes. Estas diferencias son relevantes para interpretar si la evolucion agregada se explica por mejoras generalizadas o por cambios concentrados en ciertos grupos.

## 6. Exploracion univariada y no respuesta

La tabla `outputs/tables/resumen_univariado.csv` resume faltantes, percentiles y posibles valores atipicos para variables clave. En ingresos, la no respuesta requiere tratamiento especifico. En este primer corte se considera no respuesta operativa de ingresos en ocupados cuando `P21` es faltante o negativo. Los ingresos iguales a cero no se imputan automaticamente, ya que pueden corresponder a ocupados sin ingreso laboral monetario declarado en el periodo.

## 7. Modelo de imputacion de ingresos

Antes del modelo de imputacion se estimaron dos modelos de regresion lineal vistos en clase. El primero es una regresion lineal simple, donde el logaritmo del ingreso real se explica unicamente por la edad. El segundo es una regresion lineal multiple, donde se incorporan edad, sexo, nivel educativo, aglomerado, anio, trimestre, rama de actividad (`PP04B_COD`) y ocupacion (`PP04D_COD`). Estos modelos permiten mostrar como mejora la capacidad explicativa cuando se pasa de una relacion bivariada a una especificacion multivariada.

| Modelo | n train | n test | MAE | RMSE | R2 log |
|---|---:|---:|---:|---:|---:|
| {linear_metrics.loc[0, 'modelo']} | {int(linear_metrics.loc[0, 'n_train'])} | {int(linear_metrics.loc[0, 'n_test'])} | {linear_metrics.loc[0, 'mae_pesos_2025t4']:.0f} | {linear_metrics.loc[0, 'rmse_pesos_2025t4']:.0f} | {linear_metrics.loc[0, 'r2_log']:.3f} |
| {linear_metrics.loc[1, 'modelo']} | {int(linear_metrics.loc[1, 'n_train'])} | {int(linear_metrics.loc[1, 'n_test'])} | {linear_metrics.loc[1, 'mae_pesos_2025t4']:.0f} | {linear_metrics.loc[1, 'rmse_pesos_2025t4']:.0f} | {linear_metrics.loc[1, 'r2_log']:.3f} |

Para la imputacion de no respuesta se ajusto ademas un modelo Ridge sobre el logaritmo del ingreso real de la ocupacion principal (`log(P21_REAL)`). Ridge mantiene una estructura lineal, pero agrega regularizacion para reducir la inestabilidad de coeficientes cuando hay muchas categorias de rama y ocupacion. El modelo se entreno con ocupados con ingreso positivo y se evaluo con una particion train/test. Luego se aplico a los ocupados con `P21` faltante o negativo para generar ingresos imputados.

Metricas principales:

| Modelo | n train | n test | MAE | RMSE | R2 log |
|---|---:|---:|---:|---:|---:|
| {model_metrics.loc[0, 'modelo']} | {int(model_metrics.loc[0, 'n_train'])} | {int(model_metrics.loc[0, 'n_test'])} | {model_metrics.loc[0, 'mae_pesos_2025t4']:.0f} | {model_metrics.loc[0, 'rmse_pesos_2025t4']:.0f} | {model_metrics.loc[0, 'r2_log']:.3f} |

La interpretacion de coeficientes debe hacerse en terminos aproximados de cambios porcentuales sobre el ingreso real. La tabla `outputs/tables/modelo_imputacion_coeficientes_top.csv` lista los efectos de mayor magnitud. La tabla `outputs/tables/ingresos_reales_imputados_trimestrales.csv` compara la mediana de ingresos original con la mediana luego de imputar no respuesta. Al tratarse de un modelo lineal regularizado, su principal ventaja es la interpretabilidad; su limite es que puede no capturar no linealidades complejas del mercado laboral.

## 8. Sintesis comparativa

La tabla `outputs/tables/resumen_inicial_final.csv` resume el primer y ultimo periodo observado para tasas laborales e ingresos reales. Como primer resultado, el trabajo muestra que ambos aglomerados atravesaron cambios laborales relevantes en el periodo, con una ruptura visible durante 2020 y una recomposicion posterior. La comparacion de ingresos reales muestra una evolucion condicionada por la alta inflacion del periodo, por lo que el ajuste por IPC es imprescindible para evitar conclusiones nominales engañosas.

## 9. Limitaciones y proximos ajustes

Esta version es una base de trabajo. Quedan tres puntos para revisar con cuidado: confirmar con el diseno de registro la codificacion exacta de no respuesta de ingresos; decidir si se usa IPC nacional o algun IPC regional/provincial para sensibilidad; y seleccionar los graficos finales para que el informe quede entre 6 y 10 paginas.
"""

    report_md = REPORT_DIR / "informe_base.md"
    report_md.write_text(md, encoding="utf-8")

    doc = Document()
    doc.add_heading("Evolucion del mercado laboral e ingresos", level=0)
    doc.add_paragraph("Neuquen-Plottier vs Mar del Plata")
    doc.add_heading("Introduccion", level=1)
    doc.add_paragraph(
        "Este documento es una primera version editable del informe. Resume tasas laborales, ingresos reales, "
        "analisis por subgrupos y un modelo inicial de imputacion de ingresos."
    )
    doc.add_heading("Graficos principales", level=1)
    for title, filename in [
        ("Tasa de actividad", "tasa_actividad.png"),
        ("Tasa de empleo", "tasa_empleo.png"),
        ("Tasa de desocupacion", "tasa_desocupacion.png"),
        ("Ingreso mediano real", "ingreso_mediano_real.png"),
        ("Desocupacion por sexo", "desocupacion_por_sexo.png"),
        ("Empleo por nivel educativo", "empleo_por_nivel_educativo_2025.png"),
    ]:
        doc.add_heading(title, level=2)
        doc.add_picture(str(OUTPUT_FIGURES / filename), width=Inches(6.2))
    doc.add_heading("Modelo de imputacion", level=1)
    doc.add_paragraph(
        f"Modelo: {model_metrics.loc[0, 'modelo']}. "
        f"MAE: {model_metrics.loc[0, 'mae_pesos_2025t4']:.0f}; "
        f"RMSE: {model_metrics.loc[0, 'rmse_pesos_2025t4']:.0f}; "
        f"R2 sobre log ingreso: {model_metrics.loc[0, 'r2_log']:.3f}."
    )
    doc.add_heading("Tablas de apoyo", level=1)
    doc.add_paragraph(
        "Ver archivos CSV en outputs/tables, incluyendo la serie de ingresos imputados."
    )
    doc.save(REPORT_DIR / "informe_base.docx")


def main() -> None:
    ensure_dirs()
    ipc_q = download_ipc()
    df = load_filtered_data()
    indicators = calculate_labor_indicators(df)
    calculate_subgroup_tables(df)
    df, _ = add_real_income(df, ipc_q)
    incomes = calculate_income_tables(df)
    univariate = calculate_univariate_summary(df)
    linear_metrics = fit_linear_income_models(df)
    model_metrics, model, features = fit_income_model(df)
    imputed_incomes = calculate_imputed_income_tables(df, model, features)
    make_plots(indicators, incomes, df)
    summary = first_last_summary(indicators, incomes)
    write_report(
        indicators,
        incomes,
        univariate,
        linear_metrics,
        model_metrics,
        summary,
        imputed_incomes,
    )

    filtered_path = ROOT / "data" / "infoProcesada" / "eph" / "eph_aglomerados_17_34.csv.gz"
    df.to_csv(filtered_path, index=False, compression="gzip")
    print("Analisis terminado")
    print(f"Base filtrada: {filtered_path}")
    print(f"Tablas: {OUTPUT_TABLES}")
    print(f"Graficos: {OUTPUT_FIGURES}")
    print(f"Informe md: {REPORT_DIR / 'informe_base.md'}")
    print(f"Informe docx: {REPORT_DIR / 'informe_base.docx'}")


if __name__ == "__main__":
    main()
